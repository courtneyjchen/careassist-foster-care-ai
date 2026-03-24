"""Generate the CareAssist class walkthrough Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(22)
style_h1.font.color.rgb = RGBColor(0x4A, 0x14, 0x8C)  # purple

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(16)
style_h2.font.color.rgb = RGBColor(0x4A, 0x14, 0x8C)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(13)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('CareAssist v4 — Class Walkthrough', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('AI-Driven Case Prioritization for Foster Care\nBerkeley · Spring 2026')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# ============================================================
# SECTION 1 — WHAT IS CAREASSIST?  (~1 min)
# ============================================================
doc.add_heading('1.  What Is CareAssist?', level=1)

doc.add_paragraph(
    'CareAssist is a full-stack web application that helps foster care social workers '
    'figure out which children on their caseload need attention most urgently — and why.'
)
doc.add_paragraph(
    'The foster care system in the U.S. serves over 400,000 children at any given time. '
    'Social workers often carry caseloads of 20–40+ children, each with complex trauma histories, '
    'medical needs, court deadlines, and permanency goals. Today, prioritization is mostly '
    'gut instinct and sticky notes. CareAssist gives workers a data-driven co-pilot.'
)
doc.add_paragraph(
    'The platform has four user roles — Social Worker, Supervisor, Foster Parent, and Aged-Out Youth — '
    'each with a completely different dashboard and experience tailored to what they need.'
)

# ============================================================
# SECTION 2 — TECH STACK  (~1 min)
# ============================================================
doc.add_heading('2.  Technology Stack', level=1)

table = doc.add_table(rows=6, cols=2, style='Light List Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ('Layer', 'Technology'),
    ('Frontend', 'Angular 17 (standalone components, TypeScript, lazy-loaded routes)'),
    ('Backend', 'FastAPI (Python, async, Pydantic schemas)'),
    ('Database', 'SQLite via SQLAlchemy async ORM (11 tables) — swappable to PostgreSQL'),
    ('ML Model', 'Weighted Ensemble: XGBoost + LightGBM + CatBoost + MLP (pure-Python scorer)'),
    ('AI Chat', 'Ollama / Llama 3.2 with smart rule-based fallback'),
]
for i, (a, b) in enumerate(cells):
    row = table.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why Angular? ').bold = True
p.add_run(
    'Foster care systems are government enterprise software. Angular\'s opinionated architecture '
    'with TypeScript, dependency injection, route guards, and strong typing makes it ideal for '
    'auditable, role-gated applications. Every route is lazy-loaded and every component is standalone — '
    'no NgModules.'
)

p = doc.add_paragraph()
p.add_run('Why FastAPI? ').bold = True
p.add_run(
    'Async by default (all DB queries are non-blocking via aiosqlite), auto-generates '
    'Swagger API docs, and Pydantic validates every request and response schema.'
)

# ============================================================
# SECTION 3 — THE ML MODEL  (~2 min)
# ============================================================
doc.add_heading('3.  The Machine Learning Model', level=1)

doc.add_heading('Training Data', level=2)
doc.add_paragraph(
    'Trained on AFCARS (Adoption and Foster Care Analysis and Reporting System) — the federal '
    'dataset mandated by the Children\'s Bureau. We used fiscal years 2020–2024: approximately '
    '5.76 million real, de-identified records covering every foster child in the U.S.'
)

doc.add_heading('What We Predict', level=2)
doc.add_paragraph(
    'Placement disruption — whether a child\'s current living arrangement will break down '
    'within 12 months (measured by an increase in the NUMPLEP field in AFCARS). '
    'The base disruption rate in the data is ~36%, meaning roughly 1 in 3 children '
    'experience a placement move each year.'
)

doc.add_heading('Ensemble Architecture', level=2)
doc.add_paragraph(
    'We use a weighted ensemble of four models, with weights optimized through '
    '50 Optuna trials and 5-fold stacking cross-validation:'
)

tbl2 = doc.add_table(rows=5, cols=3, style='Light List Accent 1')
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
model_rows = [
    ('Model', 'Weight', 'Key Hyperparameters'),
    ('XGBoost', '35%', '822 rounds, depth 13, lr 0.046'),
    ('LightGBM', '35%', '1,911 rounds, depth 13, 377 leaves'),
    ('CatBoost', '20%', '1,387 iterations, depth 12'),
    ('MLP', '10%', 'scikit-learn neural network'),
]
for i, (a, b, c) in enumerate(model_rows):
    row = tbl2.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('Performance', level=2)
perf = [
    'ROC-AUC: 0.9205 — the model correctly ranks a true disruption above a non-disruption 92% of the time',
    'Average Precision: 0.8615',
    'F1 Score: 0.784',
    'Recall: 92% — we optimized for high recall because missing a child at risk is worse than a false alarm',
]
for item in perf:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Features (65 total)', level=2)
doc.add_paragraph(
    '20 baseline features from AFCARS + 45 engineered features including: '
    'age², disability severity composite, abuse severity composite, '
    'substance abuse score, LOS ratio, interaction terms (age × disability, removals × abuse), '
    'one-hot placement types, permanency goals, and removal reason flags.'
)

p = doc.add_paragraph()
p.add_run('Top 5 by importance: ').bold = True
p.add_run(
    '(1) Length of stay in latest removal 13%, '
    '(2) Length of stay in current setting 12.5%, '
    '(3) Placement type 11%, '
    '(4) Total prior placements 5.2%, '
    '(5) Age at removal 4.8%.'
)

# ============================================================
# SECTION 4 — SHAP EXPLAINABILITY  (~1 min)
# ============================================================
doc.add_heading('4.  SHAP Explainability', level=1)
doc.add_paragraph(
    'A risk score alone isn\'t enough — a worker needs to know why a child is flagged. '
    'SHAP (SHapley Additive exPlanations) decomposes each prediction into per-feature contributions. '
    'Each feature gets a value: positive = increases risk (red), negative = protective (green). '
    'They sum from the baseline population risk of 36% to the child\'s predicted score.'
)
doc.add_paragraph(
    'In the app, clicking "Explain This Score in Detail" shows a horizontal bar chart of the top '
    'contributing features with human-readable labels, the child\'s actual values, and the '
    'direction and magnitude of each contribution. This is critical for trust — workers can '
    'verify the model\'s reasoning against their own clinical knowledge.'
)

# ============================================================
# SECTION 5 — APP WALKTHROUGH  (~5–7 min)
# ============================================================
doc.add_heading('5.  App Walkthrough', level=1)

# -- Login --
doc.add_heading('Login Page', level=2)
doc.add_paragraph(
    'Four role cards: Social Worker, Supervisor, Foster Parent, Aged-Out Youth. '
    'Click a card to log in (demo mode — production would use SSO/email-password). '
    'The page has a playful, child-friendly design with an animated sun cursor and '
    'crayon-drawing background — intentional warmth for software that\'s about children.'
)

# -- Social Worker --
doc.add_heading('Social Worker — Jessica Hawkins', level=2)
p = doc.add_paragraph()
p.add_run('Dashboard: ').bold = True
p.add_run(
    'Command center with 4 stat cards (Active Cases, Flagged Cases, Pending Reviews, '
    'Avg. Permanency Months), a Flagged Cases section with ML-generated alert cards, '
    'a sortable case table, and a slide-in side panel when you click a case.'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Side Panel: ').bold = True
p.add_run(
    'Shows urgency score bar, active flags, case details, plus two buttons: '
    '"Explain This Score in Detail" (opens the full case page with SHAP auto-expanded) '
    'and "Ask AI About This Case" (opens the AI assistant pre-filled with a question).'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Case Detail Page (4 tabs): ').bold = True
doc.add_paragraph(
    'Overview — Risk gauge ring (circular SVG), 6-month risk score trend chart, '
    'case info grid, SHAP explainability panel (horizontal bar chart of feature contributions), '
    'active flags with severity/confidence/recommendations, '
    'and a sibling linkage map showing connected siblings across the system.',
    style='List Bullet'
)
doc.add_paragraph(
    'Timeline — Vertical placement history timeline, color-coded by event type '
    '(placements, flags, court hearings, visits, medical). Shows provider names and dates.',
    style='List Bullet'
)
doc.add_paragraph(
    'Family — Visual family tree (parents → child → extended family). Each node shows '
    'name, relationship, phone (workers only), and safe/restricted contact badge. '
    'Workers can add new family members.',
    style='List Bullet'
)
doc.add_paragraph(
    'Notes — Shared notes (foster parent ↔ worker conversation thread with pinning) '
    'and case notes (official visit/court/general documentation).',
    style='List Bullet'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('AI Assistant: ').bold = True
p.add_run(
    'Full chat interface with quick-action suggestions. It queries the real database — '
    'ask "Which cases have the highest risk?" and it returns actual data. '
    'Uses Ollama/Llama 3.2 when available, falls back to a smart rule-based engine for demo. '
    'Also accessible via sidebar (social workers and supervisors only).'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Other pages: ').bold = True
p.add_run(
    'Messages (two-pane chat with judges, doctors, foster parents), '
    'Notifications (typed alerts: risk alerts, court reminders, flag triggers), '
    'Calendar (month grid with color-coded hearings, visits, reviews, medical appointments), '
    'Reports (per-child reports + monthly caseload summary), Files.'
)

# -- Supervisor --
doc.add_heading('Supervisor — James Chen', level=2)
doc.add_paragraph(
    'Sees a Team Overview dashboard instead of a personal caseload. '
    'Stats: total social workers, total cases, flagged cases, avg cases/worker. '
    'Expandable worker cards for each social worker (Jessica Hawkins, Priya Patel, Marcus Williams) '
    'showing their metrics, a color-coded workload bar, and a full case table when expanded. '
    'Supervisors can click into any case for the full detail page including SHAP. '
    'Also has weekly check-in schedule, team-level Reports, and AI Assistant access.'
)

# -- Foster Parent --
doc.add_heading('Foster Parent — Maria Garcia', level=2)
doc.add_paragraph(
    'Sees the Family Portal — no risk scores, no ML flags, no SHAP. '
    'Foster parents are caregivers, not case managers. They see:'
)
doc.add_paragraph('Children cards with name, age, placement type, need tags (medical/behavioral/disability)', style='List Bullet')
doc.add_paragraph('Upload Document and Message Social Worker buttons per child', style='List Bullet')
doc.add_paragraph('Messages with their social worker, pediatrician, school counselor, support groups', style='List Bullet')
doc.add_paragraph('Notifications (visit reminders, new documents, worker messages)', style='List Bullet')
doc.add_paragraph('Calendar (pediatrician visits, therapy, parent-teacher conferences)', style='List Bullet')
doc.add_paragraph('Resources directory: Training, Support Groups, Financial (stipend info), Health (Medicaid), Legal (foster parent rights)', style='List Bullet')

# -- Youth --
doc.add_heading('Aged-Out Youth — Jordan Davis', level=2)
doc.add_paragraph(
    'The Youth Portal is designed for self-advocacy. Jordan aged out of care and needs '
    'access to his own history. He sees:'
)
doc.add_paragraph('My Portal — quick stats, quick-access cards for Medical History, School Records, Support, Resources', style='List Bullet')
doc.add_paragraph('My Records — unified timeline of his Medical History, School Records, and Placement History with dates and providers', style='List Bullet')
doc.add_paragraph('Resources — Transitional Housing, Job Training, Chafee Education Grants, Tuition Waivers, Extended Medicaid, Mentorship, Legal Aid', style='List Bullet')
doc.add_paragraph('Messages — stay in touch with former case worker, life skills mentor, housing services', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why this matters: ').italic = True
p.add_run(
    'When youth age out at 18, they often lose access to their own records. '
    'CareAssist gives them a portable, consolidated record of their entire history.'
)

# ============================================================
# SECTION 6 — WHAT'S NEXT / FUTURE WORK  (~1 min)
# ============================================================
doc.add_heading('6.  Future Work', level=1)

future_items = [
    ('Sibling Linkage with Live Data', 'The sibling map currently uses seed data. In production it would cross-reference AFCARS sibling fields to find siblings across workers and counties — a real gap in the current system.'),
    ('Real LLM Integration', 'Full Ollama/Llama 3.2 conversational AI for policy questions, intervention suggestions, and court report drafting.'),
    ('S3 Document Storage', 'The backend S3 service is built but not yet connected to a live AWS bucket for medical records, court docs, and IEPs.'),
    ('Real-Time Scoring', 'Currently scores are computed at seed time. In production, scores would recalculate whenever case data changes.'),
    ('HIPAA Compliance', 'Production would need SSO/SAML auth, API-level role enforcement, audit logging, encryption at rest, and HIPAA BAA.'),
]
for title_text, desc in future_items:
    p = doc.add_paragraph()
    p.add_run(title_text + ' — ').bold = True
    p.add_run(desc)

# ============================================================
# SECTION 7 — QUICK REFERENCE
# ============================================================
doc.add_heading('7.  Quick Reference', level=1)

doc.add_heading('Demo Logins', level=2)
login_table = doc.add_table(rows=5, cols=3, style='Light List Accent 1')
login_table.alignment = WD_TABLE_ALIGNMENT.CENTER
login_rows = [
    ('Role', 'Name', 'Email'),
    ('Social Worker', 'Jessica Hawkins', 'jessica.hawkins@careassist.org'),
    ('Supervisor', 'James Chen', 'james.chen@careassist.org'),
    ('Foster Parent', 'Maria Garcia', 'maria.garcia@careassist.org'),
    ('Aged-Out Youth', 'Jordan Davis', 'jordan.davis@careassist.org'),
]
for i, (a, b, c) in enumerate(login_rows):
    row = login_table.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph('Password for all accounts: demo1234')

doc.add_heading('Key Numbers', level=2)
nums_table = doc.add_table(rows=8, cols=2, style='Light List Accent 1')
nums_table.alignment = WD_TABLE_ALIGNMENT.CENTER
nums_rows = [
    ('Metric', 'Value'),
    ('Training Data', '5.76 M records (AFCARS FY 2020–2024)'),
    ('Model', 'Weighted Ensemble (XGBoost + LightGBM + CatBoost + MLP)'),
    ('ROC-AUC', '0.9205'),
    ('Recall', '92%'),
    ('Features', '65 (20 baseline + 45 engineered)'),
    ('Seed Data', '18 children, 6 users, 3 social workers'),
    ('Base Disruption Rate', '~36%'),
]
for i, (a, b) in enumerate(nums_rows):
    row = nums_table.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_heading('Good Cases to Show', level=2)
cases_table = doc.add_table(rows=4, cols=3, style='Light List Accent 1')
cases_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cases_rows = [
    ('Child', 'Risk', 'Why'),
    ('Aisha Williams', 'High', '30 mo in care, 5 placements, group home, "Permanency Delay" critical flag — best SHAP demo'),
    ('Maya Johnson', 'High', '3 placements, behavioral escalation, neglect — good for showing flags + family tree'),
    ('Zoe Brown', 'Low', '2 months, kinship care, 0 prior placements — shows protective factors in SHAP'),
]
for i, (a, b, c) in enumerate(cases_rows):
    row = cases_table.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# ============================================================
# Save
# ============================================================
out_path = os.path.expanduser(r'~\Downloads\CareAssist_Class_Walkthrough.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
